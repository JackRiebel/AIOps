"""DOCX document parser.

Extracts text from Word documents with structure preservation,
including headings, paragraphs, tables, and lists.
"""

import io
import logging
from typing import List, Optional, BinaryIO
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class DocxSection:
    """Represents a section of a Word document."""
    heading: Optional[str]
    heading_level: int
    content: str
    tables: List[str] = field(default_factory=list)


@dataclass
class DocxDocument:
    """Represents an extracted Word document."""
    sections: List[DocxSection]
    metadata: dict
    raw_text: str

    @property
    def as_markdown(self) -> str:
        """Convert to markdown format with headers."""
        parts = []
        for section in self.sections:
            if section.heading:
                # Use appropriate header level
                prefix = "#" * min(section.heading_level, 6)
                parts.append(f"{prefix} {section.heading}")
            if section.content:
                parts.append(section.content)
            for table in section.tables:
                parts.append(table)
        return "\n\n".join(parts)


class DocxParser:
    """Parser for Word documents (DOCX).

    Uses python-docx for extraction with structure preservation.
    """

    def __init__(self):
        self._docx_available = False
        try:
            import docx
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX parsing will be unavailable.")

    def parse(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        file_stream: Optional[BinaryIO] = None,
    ) -> DocxDocument:
        """Parse a DOCX file and extract text with structure.

        Args:
            file_path: Path to DOCX file
            file_content: Raw bytes of DOCX content
            file_stream: File-like object containing DOCX

        Returns:
            DocxDocument with extracted text and metadata

        Raises:
            ValueError: If no valid input provided or python-docx not available
        """
        if not self._docx_available:
            raise ValueError("python-docx library is required for DOCX parsing. Install with: pip install python-docx")

        import docx

        # Load document
        if file_path:
            doc = docx.Document(file_path)
        elif file_content:
            doc = docx.Document(io.BytesIO(file_content))
        elif file_stream:
            doc = docx.Document(file_stream)
        else:
            raise ValueError("Must provide file_path, file_content, or file_stream")

        # Extract metadata
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
            }
            metadata = {k: v for k, v in metadata.items() if v}

        # Extract content with structure
        sections = []
        current_section = DocxSection(heading=None, heading_level=0, content="", tables=[])
        raw_text_parts = []

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                para = docx.text.paragraph.Paragraph(element, doc)
                text = para.text.strip()

                if not text:
                    continue

                raw_text_parts.append(text)

                # Check if this is a heading
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    # Save current section
                    if current_section.content or current_section.heading:
                        sections.append(current_section)

                    # Determine heading level
                    try:
                        level = int(style_name.replace("Heading ", "").replace("Heading", "1"))
                    except ValueError:
                        level = 1

                    current_section = DocxSection(
                        heading=text,
                        heading_level=level,
                        content="",
                        tables=[]
                    )
                else:
                    # Regular paragraph
                    if current_section.content:
                        current_section.content += "\n\n" + text
                    else:
                        current_section.content = text

            # Handle tables
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, doc)
                table_md = self._table_to_markdown(table)
                if table_md:
                    current_section.tables.append(table_md)
                    raw_text_parts.append(table_md)

        # Don't forget last section
        if current_section.content or current_section.heading or current_section.tables:
            sections.append(current_section)

        return DocxDocument(
            sections=sections,
            metadata=metadata,
            raw_text="\n\n".join(raw_text_parts)
        )

    def parse_to_text(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        as_markdown: bool = True,
    ) -> str:
        """Parse DOCX and return plain text or markdown.

        Args:
            file_path: Path to DOCX file
            file_content: Raw bytes of DOCX content
            as_markdown: If True, format with markdown headers

        Returns:
            Extracted text as string
        """
        doc = self.parse(file_path=file_path, file_content=file_content)

        if as_markdown:
            return doc.as_markdown
        else:
            return doc.raw_text

    def _table_to_markdown(self, table) -> str:
        """Convert a Word table to markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if not rows:
            return ""

        # Add header separator after first row
        if len(rows) >= 1:
            num_cols = rows[0].count("|") - 1
            separator = "| " + " | ".join(["---"] * num_cols) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)

    def extract_headers(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
    ) -> List[dict]:
        """Extract just the document outline (headers).

        Args:
            file_path: Path to DOCX file
            file_content: Raw bytes of DOCX content

        Returns:
            List of headers with text and level
        """
        doc = self.parse(file_path=file_path, file_content=file_content)

        headers = []
        for section in doc.sections:
            if section.heading:
                headers.append({
                    "text": section.heading,
                    "level": section.heading_level
                })

        return headers


# Singleton instance
_docx_parser: Optional[DocxParser] = None


def get_docx_parser() -> DocxParser:
    """Get or create the DocxParser singleton."""
    global _docx_parser
    if _docx_parser is None:
        _docx_parser = DocxParser()
    return _docx_parser
